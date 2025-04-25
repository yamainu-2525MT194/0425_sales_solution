
# 完全自動日報生成ツール（ChatGPT API連携）

import openai
from docx import Document
import string
import time
from datetime import datetime

# --- ChatGPT APIキーをここに設定（必要） ---

# 名前生成関数（A〜CVの100人分）
def generate_names(n):
    names = []
    for i in range(n):
        if i < 26:
            names.append(string.ascii_uppercase[i])
        else:
            first = string.ascii_uppercase[(i - 26) // 26]
            second = string.ascii_uppercase[(i - 26) % 26]
            names.append(first + second)
    return names

# ChatGPT（gpt-3.5-turbo）で1人分の日報を完全自動生成（数値・文字込み）
def generate_report_content(name, today):
    prompt = f"""
以下のフォーマットに従って、実在しない内容で自然な日本語の日報を生成してください。
数値（X部分）は0〜5のランダムな件数、○○部分は架空の名前や内容を自動生成してください。

名前：{name}
日付：{today}

【成果】
■数値
・オファー数：X件
・面談獲得数：X件
・提案数（有効）：X件
・提案数（有効・無効・不明）：X件
・配信数：XXX件　本日X名追加
・面談実施数：X件
　対象エンジニア名：○○さん
　エンジニア所感：○○○○
　顧客所感：○○○○
　確度：約XX％

■ヒアリング
・○○さん：○○○○

■打ち合わせ
・○○企業 ○○様：○○○○

■アシスト
・○○○○

【振り返り】
■良かった点：
・○○○○

■課題：
・○○○○

■対策：
・○○○○

必ずフォーマットに従って構成し、数値と文字は実際に自動生成してください。
"""
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.8
    )
    return response['choices'][0]['message']['content']

# Wordファイルとして保存
def create_word_report(filename="Auto_Generated_Daily_Report_Final.docx"):
    doc = Document()
    today = datetime.now().strftime('%Y年%m月%d日')
    doc.add_heading(f'自動生成：日報報告資料（{today} / gpt-3.5-turbo）', 0)

    names = generate_names(50)
    for i, name in enumerate(names):
        try:
            print(f"{i+1}/50 ▶ 名前：{name} を生成中...")
            report = generate_report_content(name, today)
            doc.add_heading(f'名前：{name}', level=1)
            for line in report.strip().split('\n'):
                doc.add_paragraph(line.strip())
            time.sleep(1)  # API制限対応
        except Exception as e:
            print(f"❌ {name} さんの生成に失敗しました: {repr(e)}")

    doc.save(filename)
    print(f"✅ Wordファイル「{filename}」を作成しました。")

# 実行
if __name__ == '__main__':
    create_word_report()
